"""DOCX parser using python-docx."""
from docx import Document as DocxDocument


def parse_docx(file_path: str) -> dict:
    """
    Extract text and tables from a DOCX file.
    Returns dict with keys: text, tables, metadata
    """
    doc = DocxDocument(file_path)
    paragraphs = []
    tables = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append({
                "text": para.text,
                "style": para.style.name if para.style else None,
                "is_heading": para.style.name.startswith("Heading") if para.style else False,
            })

    # Extract tables
    for table_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            tables.append({"index": table_idx, "rows": rows})

    full_text = "\n".join(p["text"] for p in paragraphs)

    return {
        "text": full_text,
        "paragraphs": paragraphs,
        "tables": tables,
        "metadata": {
            "author": doc.core_properties.author,
            "title": doc.core_properties.title,
        },
    }
